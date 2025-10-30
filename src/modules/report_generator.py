"""
Professional Report Generator Module
Creates Word documents and PDF reports from PDF content
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
from typing import List, Dict, Optional
import os


class ReportGenerator:
    """Generate professional Word and PDF reports"""
    
    def __init__(self, config: dict = None):
        """
        Initialize report generator
        
        Args:
            config: Configuration dictionary
        """
        self.config = config or self.get_default_config()
        self.doc = None
        
    def get_default_config(self) -> dict:
        """Get default configuration"""
        return {
            'page_size': 'A4',
            'title_font_size': 24,
            'heading_font_size': 16,
            'content_font_size': 11,
            'title_color': (0, 51, 102),
            'margins': {'top': 1.0, 'bottom': 1.0, 'left': 1.0, 'right': 1.0}
        }
    
    def create_word_document(self, title: str = "Report") -> Document:
        """
        Create a new Word document
        
        Args:
            title: Document title
            
        Returns:
            Document object
        """
        self.doc = Document()
        
        # Set margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(self.config['margins']['top'])
            section.bottom_margin = Inches(self.config['margins']['bottom'])
            section.left_margin = Inches(self.config['margins']['left'])
            section.right_margin = Inches(self.config['margins']['right'])
        
        return self.doc
    
    def add_title(self, title: str):
        """
        Add title to document
        
        Args:
            title: Document title
        """
        heading = self.doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Format title
        run = heading.runs[0]
        run.font.size = Pt(self.config.get('title_font_size', 24))
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.config.get('title_color', (0, 51, 102)))
    
    def add_heading(self, heading_text: str, level: int = 1):
        """
        Add heading to document
        
        Args:
            heading_text: Heading text
            level: Heading level (1-3)
        """
        heading = self.doc.add_heading(heading_text, level=level)
        run = heading.runs[0]
        run.font.size = Pt(self.config.get('heading_font_size', 16))
        run.font.color.rgb = RGBColor(*self.config.get('title_color', (0, 51, 102)))
    
    def add_paragraph(self, text: str):
        """
        Add paragraph to document
        
        Args:
            text: Paragraph text
        """
        p = self.doc.add_paragraph(text)
        
        # Format paragraph
        for run in p.runs:
            run.font.size = Pt(self.config.get('content_font_size', 11))
            run.font.name = 'Arial'
        
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def add_bullet_list(self, items: List[str]):
        """
        Add bullet list to document
        
        Args:
            items: List of items
        """
        for item in items:
            p = self.doc.add_paragraph(item, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(self.config.get('content_font_size', 11))
    
    def add_page_break(self):
        """Add page break"""
        self.doc.add_page_break()
    
    def generate_word_report(self, text_chunks: List[str], title: str,
                            subtitle: str = None, output_path: str = None) -> str:
        """
        Generate Word report from text chunks
        
        Args:
            text_chunks: List of text chunks
            title: Report title
            subtitle: Report subtitle
            output_path: Output file path
            
        Returns:
            Path to generated report
        """
        # Create document
        self.create_word_document(title)
        
        # Add title
        self.add_title(title)
        
        # Add subtitle if provided
        if subtitle:
            p = self.doc.add_paragraph(subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.italic = True
        
        # Add spacing
        self.doc.add_paragraph()
        
        # Add content sections
        for i, chunk in enumerate(text_chunks, 1):
            lines = chunk.strip().split('\n')
            
            if lines:
                # First line as heading
                section_title = lines[0][:100] if len(lines[0]) > 100 else lines[0]
                self.add_heading(f"Section {i}: {section_title}", level=2)
                
                # Rest as content
                if len(lines) > 1:
                    content = '\n'.join(lines[1:])
                    self.add_paragraph(content)
                
                # Add spacing between sections
                if i < len(text_chunks):
                    self.doc.add_paragraph()
        
        # Save document
        if output_path is None:
            output_dir = os.path.join('output', 'reports')
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, f"{title.replace(' ', '_')}.docx")
        
        self.doc.save(output_path)
        return output_path
    
    def generate_pdf_report(self, text_chunks: List[str], title: str,
                           subtitle: str = None, output_path: str = None) -> str:
        """
        Generate PDF report from text chunks
        
        Args:
            text_chunks: List of text chunks
            title: Report title
            subtitle: Report subtitle
            output_path: Output file path
            
        Returns:
            Path to generated PDF report
        """
        # Set output path
        if output_path is None:
            output_dir = os.path.join('output', 'reports')
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, f"{title.replace(' ', '_')}.pdf")
        
        # Create PDF
        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            topMargin=inch,
            bottomMargin=inch,
            leftMargin=inch,
            rightMargin=inch
        )
        
        # Container for PDF elements
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        
        # Custom title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#003366'),
            spaceAfter=30,
            alignment=TA_CENTER,
            bold=True
        )
        
        # Custom heading style
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#003366'),
            spaceAfter=12,
            spaceBefore=12,
            bold=True
        )
        
        # Custom body style
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        )
        
        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Add subtitle if provided
        if subtitle:
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Normal'],
                fontSize=14,
                alignment=TA_CENTER,
                italic=True,
                spaceAfter=30
            )
            story.append(Paragraph(subtitle, subtitle_style))
        
        # Add content sections
        for i, chunk in enumerate(text_chunks, 1):
            lines = chunk.strip().split('\n')
            
            if lines:
                # First line as heading
                section_title = lines[0][:100] if len(lines[0]) > 100 else lines[0]
                story.append(Paragraph(f"Section {i}: {section_title}", heading_style))
                
                # Rest as content
                if len(lines) > 1:
                    content = '<br/>'.join(lines[1:])
                    story.append(Paragraph(content, body_style))
                
                # Add spacing between sections
                if i < len(text_chunks):
                    story.append(Spacer(1, 0.2 * inch))
        
        # Build PDF
        doc.build(story)
        return output_path
